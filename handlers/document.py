"""Handler for uploaded documents (DOCX, PDF)."""

import os
import time as _time
from typing import Optional

import io

from aiogram import Bot, Router
from aiogram.types import BufferedInputFile, Message

from config import settings
from database.crud import save_search_history
from services.docx_parser import extract_text_from_docx
from services.excel_generator import generate_report
from services.matcher import find_matching_models
from services.table_parser import parse_requirements_from_tables, parse_inline_descriptions
from utils.logger import logger


def _extract_text_from_docx(file_path: str) -> str:
    """Extract plain text from DOCX paragraphs for LLM input."""
    try:
        from docx import Document
        doc = Document(file_path)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_texts:
                    parts.append(" | ".join(row_texts))
        return "\n".join(parts)
    except Exception as e:
        logger.warning(f"_extract_text_from_docx failed: {e}")
        return ""

router = Router()

TEMP_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "temp_files")


async def _safe_edit(msg, text: str) -> None:
    """Edit message text, ignoring 'message is not modified' errors."""
    try:
        await msg.edit_text(text)
    except Exception as e:
        if "message is not modified" not in str(e):
            raise


async def extract_from_pdf(file_path: str) -> Optional[str]:
    """Extract text and tables from a PDF file. Returns None if extraction fails."""
    try:
        import pdfplumber
    except ImportError:
        logger.error("pdfplumber not installed — PDF support unavailable")
        return None

    try:
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                tables = page.extract_tables()
                for table in tables:
                    if not table:
                        continue
                    for row in table:
                        if row:
                            text_parts.append("\t".join(
                                cell if cell is not None else "" for cell in row
                            ))
                text = page.extract_text()
                if text:
                    text_parts.append(text)

        if not text_parts:
            return None
        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        return None


async def _parse_pdf_inline(file_path: str) -> Optional[dict]:
    """
    Parse a PDF by extracting text, then trying inline pattern matching.
    Returns requirements dict or None.
    """
    pdf_text = await extract_from_pdf(file_path)
    if not pdf_text:
        return None

    try:
        # Build a minimal requirements dict from text patterns
        from services.table_parser import _parse_inline_block, parse_value
        import re

        items = []
        current_specs = {}
        current_item_name = None
        current_model_name = None
        current_category = None

        for line in pdf_text.split('\n'):
            line = line.strip()
            if not line:
                continue

            m = re.match(r'^([^:]{2,60}):\s*(.+)$', line)
            if not m:
                if ';' in line:
                    extra = _parse_inline_block(line)
                    current_specs.update(extra)
                continue

            key = m.group(1).strip()
            value_raw = m.group(2).strip()
            key_lower = key.lower()

            if re.search(r'наименование\s*(товара|изделия|позиции|оборудования)', key_lower):
                if current_specs:
                    from services.table_parser import _extract_model_name_from_text
                    cat = current_category
                    if cat is None and current_item_name:
                        il = current_item_name.lower()
                        if "коммутатор" in il or "switch" in il:
                            cat = "Коммутаторы"
                        elif "маршрутизатор" in il or "router" in il or "шлюз" in il:
                            cat = "Маршрутизаторы"
                    items.append({
                        "item_name": current_item_name or f"Позиция {len(items)+1}",
                        "quantity": None,
                        "model_name": current_model_name,
                        "category": cat,
                        "required_specs": dict(current_specs),
                    })
                current_specs = {}
                current_item_name = value_raw
                from services.table_parser import _extract_model_name_from_text
                current_model_name = _extract_model_name_from_text(value_raw)
                current_category = None
                continue

            if re.search(r'(категория|тип\s*устройства)', key_lower):
                vl = value_raw.lower()
                if "коммутатор" in vl or "switch" in vl:
                    current_category = "Коммутаторы"
                elif "маршрутизатор" in vl or "router" in vl or "шлюз" in vl:
                    current_category = "Маршрутизаторы"
                continue

            if re.search(r'количество\s*(единиц|шт\.?|штук)', key_lower):
                continue

            parsed = parse_value(value_raw)
            if parsed is not None:
                current_specs[key] = parsed

        if current_specs:
            cat = current_category
            if cat is None and current_item_name:
                il = (current_item_name or "").lower()
                if "коммутатор" in il or "switch" in il:
                    cat = "Коммутаторы"
                elif "маршрутизатор" in il or "router" in il:
                    cat = "Маршрутизаторы"
            items.append({
                "item_name": current_item_name or "Позиция 1",
                "quantity": None,
                "model_name": current_model_name,
                "category": cat,
                "required_specs": dict(current_specs),
            })

        if items:
            return {"items": items}
        return None
    except Exception as e:
        logger.error(f"PDF inline parsing error: {e}")
        return None


@router.message(lambda m: m.document is not None)
async def handle_document(message: Message, bot: Bot) -> None:
    """Download and process an uploaded document (DOCX or PDF)."""
    doc = message.document
    file_name = doc.file_name or "unknown"
    user_id = message.from_user.id
    logger.info(f"Document received from {user_id}: {file_name} ({doc.file_size} bytes)")

    fname_lower = file_name.lower()
    is_pdf = fname_lower.endswith(".pdf")
    is_docx = fname_lower.endswith(".docx")

    if not is_docx and not is_pdf:
        await message.answer(
            "Неподдерживаемый формат файла.\n"
            "Пожалуйста, отправьте файл в формате DOCX или PDF."
        )
        return

    if doc.file_size and doc.file_size > 20 * 1024 * 1024:
        await message.answer("Файл слишком большой (макс. 20 МБ).")
        return

    status_msg = await message.answer("Файл получен. Начинаю анализ...")

    os.makedirs(TEMP_DIR, exist_ok=True)
    file_path = os.path.join(TEMP_DIR, f"{user_id}_{file_name}")
    excel_path = None

    try:
        # Download file
        file = await bot.get_file(doc.file_id)
        await bot.download_file(file.file_path, file_path)
        logger.info(f"File downloaded to {file_path}")
        _start_time = _time.time()

        await _safe_edit(status_msg,"Анализирую структуру документа...")

        requirements = None

        # ── DOCX path ──
        if is_docx:
            # Step 0: LLM parser (primary, if enabled)
            if settings.llm_parsing_enabled and settings.openai_api_key:
                logger.info("Attempting LLM-based parsing...")
                await _safe_edit(status_msg,"Анализирую документ с помощью ИИ...")
                from services.llm_parser import parse_tz_with_llm
                doc_text = _extract_text_from_docx(file_path)
                if doc_text:
                    requirements = await parse_tz_with_llm(doc_text)
                    if requirements and requirements.get("items"):
                        items = requirements.get("items", [])
                        logger.info(f"✓ LLM parser succeeded: {len(items)} items extracted")
                        await _safe_edit(status_msg,
                            f"✓ ИИ извлёк требования\n"
                            f"Извлечено позиций: {len(items)}"
                        )

            # Step 1: table-based parsing (fallback)
            if not requirements:
                logger.info("Attempting table-based parsing...")
                await _safe_edit(status_msg,"Анализирую структуру документа...")
                requirements = parse_requirements_from_tables(file_path)

                if requirements and requirements.get("items"):
                    items = requirements.get("items", [])
                    logger.info(f"✓ Table parser succeeded: {len(items)} items extracted")
                    await _safe_edit(status_msg,
                        f"✓ Обнаружена структурированная таблица\n"
                        f"Извлечено позиций: {len(items)}"
                    )

            # Step 2: inline parser (last resort)
            if not requirements:
                logger.info("Table parser found nothing, trying inline parser...")
                await _safe_edit(status_msg,"Таблица не найдена, анализирую описания...")
                requirements = parse_inline_descriptions(file_path)

                if requirements and requirements.get("items"):
                    items = requirements.get("items", [])
                    logger.info(f"✓ Inline parser succeeded: {len(items)} items extracted")
                    await _safe_edit(status_msg,
                        f"✓ Обнаружены описания оборудования\n"
                        f"Извлечено позиций: {len(items)}"
                    )
                else:
                    logger.info("All parsers failed — returning error to user")
                    await _safe_edit(status_msg,
                        "Не удалось распознать структуру документа.\n\n"
                        "Бот поддерживает:\n"
                        "1. Таблица с колонками: Наименование | № | Характеристика | Значение | Ед.изм.\n"
                        "2. Описание вида «Характеристика: значение» (по одной на строку или через «;»)\n"
                        "3. Произвольный текст ТЗ (при настроенном OPENAI_API_KEY)\n\n"
                        "Если документ соответствует одному из форматов — обратитесь к администратору."
                    )
                    return

        # ── PDF path ──
        elif is_pdf:
            logger.info("Extracting text from PDF...")
            await _safe_edit(status_msg,"Извлекаю текст из PDF...")
            requirements = await _parse_pdf_inline(file_path)

            if not (requirements and requirements.get("items")):
                await _safe_edit(status_msg,
                    "Не удалось распознать структуру PDF.\n"
                    "Пожалуйста, отправьте файл в формате DOCX."
                )
                return

            items = requirements.get("items", [])
            logger.info(f"✓ PDF inline parser: {len(items)} items extracted")
            await _safe_edit(status_msg,
                f"✓ PDF обработан\nИзвлечено позиций: {len(items)}"
            )

        items = (requirements or {}).get("items", [])

        if not items:
            await _safe_edit(status_msg,
                "Не удалось извлечь требования к оборудованию из документа.\n"
                "Убедитесь, что файл содержит техническое задание с характеристиками."
            )
            return

        # Summary of extracted items
        summary_lines = [f"Извлечено позиций оборудования: {len(items)}\n"]
        for i, item in enumerate(items, 1):
            name = item.get("item_name") or item.get("model_name") or "Без названия"
            category = item.get("category") or "—"
            specs_count = len({k: v for k, v in item.get("required_specs", {}).items() if not k.startswith("__")})
            model = item.get("model_name")
            model_str = f" (модель: {model})" if model else ""
            summary_lines.append(
                f"{i}. {name}{model_str}\n   Категория: {category}, характеристик: {specs_count}"
            )

        summary_text = "\n".join(summary_lines)

        await _safe_edit(status_msg,
            f"{summary_text}\n\n"
            "Этап 2/3: Сопоставление с базой данных..."
        )

        match_results = await find_matching_models(requirements)
        match_summary = match_results.get("summary", {})

        result_lines = [
            f"\nРезультаты сопоставления:",
            f"Найдено моделей: {match_summary.get('total_models_found', 0)}",
            f"Идеальные совпадения: {match_summary.get('ideal_matches', 0)}",
            f"Частичные совпадения: {match_summary.get('partial_matches', 0)}",
        ]

        for idx, result in enumerate(match_results.get("results", []), 1):
            req = result["requirement"]
            matches = result["matches"]
            ideal = matches.get("ideal", [])
            partial = matches.get("partial", [])
            req_name = req.get("item_name") or req.get("model_name") or f"Позиция {idx}"

            if ideal:
                top = ideal[0]
                ver = top.get("version") or top.get("source_filename", "")
                result_lines.append(
                    f"\n{idx}. {req_name}:\n"
                    f"   ✅ {top['model_name']} ({ver}) — 100%"
                )
            elif partial:
                top = partial[0]
                ver = top.get("version") or top.get("source_filename", "")
                result_lines.append(
                    f"\n{idx}. {req_name}:\n"
                    f"   ⚠️ {top['model_name']} ({ver}) — {top['match_percentage']}%"
                )
            else:
                result_lines.append(f"\n{idx}. {req_name}:\n   ❌ Подходящих моделей не найдено")

        match_text = "\n".join(result_lines)

        await _safe_edit(status_msg,
            f"{summary_text}\n{match_text}\n\n"
            "Этап 3/3: Генерация Excel отчета..."
        )

        excel_path = generate_report(
            requirements=requirements,
            match_results=match_results,
            output_dir=TEMP_DIR,
            min_percentage=float(settings.match_threshold),
            filename=file_name,
            processing_time=_time.time() - _start_time,
        )

        try:
            await save_search_history(
                user_id=user_id,
                docx_filename=file_name,
                requirements=requirements,
                results_summary=match_summary,
            )
        except Exception as e:
            logger.error(f"Failed to save search history: {e}")

        # Read file into memory before deletion so the upload doesn't depend on the temp file
        excel_filename = os.path.basename(excel_path)
        with open(excel_path, "rb") as _f:
            excel_bytes = _f.read()
        os.remove(excel_path)
        excel_path = None  # mark as already removed so finally block skips it
        logger.debug(f"Temp Excel loaded into memory and removed: {excel_filename}")

        excel_file = BufferedInputFile(excel_bytes, filename=excel_filename)
        await message.answer_document(
            document=excel_file,
            caption=(
                f"Отчет готов!\n\n"
                f"{summary_text}\n{match_text}\n\n"
                f"📊 Детальное сравнение — в приложенном Excel файле."
            ),
        )

        await status_msg.delete()

        logger.info(
            f"Document processed for user {user_id}: {len(items)} items, "
            f"{match_summary.get('total_models_found', 0)} models found"
        )

    except ValueError as e:
        logger.error(f"Document parsing error for user {user_id}: {e}")
        await _safe_edit(status_msg,f"Ошибка при обработке документа:\n{e}")
    except Exception as e:
        logger.error(f"Unexpected error processing document for user {user_id}: {e}", exc_info=True)
        await _safe_edit(status_msg,
            "Произошла ошибка при обработке документа.\n"
            "Попробуйте позже или обратитесь к администратору."
        )
    finally:
        if os.path.exists(file_path):
            os.remove(file_path)
            logger.debug(f"Temp file removed: {file_path}")
        if excel_path and os.path.exists(excel_path):
            os.remove(excel_path)
            logger.debug(f"Temp Excel removed: {excel_path}")
