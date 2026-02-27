"""
Парсеры требований из документов ТЗ (DOCX).

1. parse_requirements_from_tables() — структурированные таблицы (5-колоночный формат)
2. parse_inline_descriptions() — строчный формат «Характеристика: значение; ...»

Оба возвращают оригинальные названия характеристик без нормализации к canonical_name.
"""

import re
from collections import defaultdict
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from utils.logger import logger


# Составное условие вида "> 512 и ≤ 1024": каждый токен — оператор + число
_compound_re = re.compile(
    r'(?:[≥≤><≠=]+|>=|<=|!=|не\s+менее|не\s+более|до)\s*\d[\d,.]*',
    re.IGNORECASE,
)



def parse_value(value_str: str, unit: str = "") -> Any:
    """
    Разбор значения характеристики из ячейки таблицы.
    Сохраняет операторы сравнения (≥, ≤ и т. д.) в возвращаемой строке.
    """
    if not value_str:
        return None

    value_str = value_str.replace('\xa0', ' ').strip()

    if value_str.lower() in ["да", "yes", "истина", "true"]:
        return True
    if value_str.lower() in ["нет", "no", "ложь", "false"]:
        return False

    # Составное условие: "> 512 и ≤ 1024" → список условий
    parts = _compound_re.findall(value_str)
    if len(parts) >= 2:
        parsed = [parse_value(p.strip()) for p in parts]
        return [p for p in parsed if p is not None]

    operator_match = re.match(r'^([≥≤><≠]=?|>=|<=|!=)\s*', value_str)
    operator = None
    if operator_match:
        raw_op = operator_match.group(1)
        op_map = {'≥': '>=', '≤': '<=', '≠': '!=', '>': '>', '<': '<', '=': '=',
                  '≥=': '>=', '≤=': '<='}
        operator = op_map.get(raw_op, raw_op)

    if not operator:
        if re.match(r'не\s+менее\b', value_str, re.IGNORECASE):
            operator = ">="
        elif re.match(r'не\s+более\b', value_str, re.IGNORECASE):
            operator = "<="
        elif re.match(r'до\s+', value_str, re.IGNORECASE):
            operator = "<="

    # Убираем префикс оператора перед извлечением числа
    value_for_numbers = re.sub(r'^[≥≤><≠=]+\s*', '', value_str)
    value_for_numbers = re.sub(
        r'^(?:не\s+менее|не\s+более|до)\s+', '', value_for_numbers, flags=re.IGNORECASE
    )

    # Сложение: "24+4" → 28
    sum_match = re.match(r'^(\d+(?:\.\d+)?)\s*\+\s*(\d+(?:\.\d+)?)$', value_for_numbers.strip())
    if sum_match:
        num_val = float(sum_match.group(1)) + float(sum_match.group(2))
        return f"{operator}{num_val}" if operator else num_val

    # Умножение: "24x4" → 96
    mult_match = re.match(r'^(\d+)\s*[xхX×]\s*(\d+)$', value_for_numbers.strip())
    if mult_match:
        num_val = int(mult_match.group(1)) * int(mult_match.group(2))
        return f"{operator}{num_val}" if operator else num_val

    numbers = re.findall(r'[\d,]+\.?\d*', value_for_numbers)
    if numbers:
        num_str = numbers[0].replace(',', '')
        try:
            num_val = int(num_str) if '.' not in num_str else float(num_str)
            # Сохраняем единицу измерения (например, "Гбит/с") — матчер применит множитель
            unit_in_str = re.search(
                r'\b([a-zA-Zа-яА-Я/]+(?:/[a-zA-Zа-яА-Я]+)?)\s*$', value_for_numbers
            )
            effective_unit = unit_in_str.group(1) if unit_in_str else ""
            if effective_unit:
                result_str = f"{num_val} {effective_unit}"
                return f"{operator}{result_str}" if operator else result_str
            return f"{operator}{num_val}" if operator else num_val
        except ValueError:
            pass

    return value_str


_ITEM_NAME_PATTERNS = [
    r'наименование\s*(товара|оборудования|изделия|позиции)',
    r'тип\s*(оборудования|устройства)',
    r'раздел',
]
_ITEM_NUMBER_PATTERNS = [
    r'№\s*п/?п',
    r'^п/?п$',
    r'номер',
    r'^\d+$',
]
_CHAR_NAME_PATTERNS = [
    r'наименование\s*характеристик',
    r'характеристик',
    r'параметр',
    r'требование',
    r'показатель',
]
_VALUE_PATTERNS = [
    r'значение\s*(характеристики|параметра)?',
    r'требуемое\s*значение',
    r'^значение$',
    r'величина',
    r'^value$',
]
_UNIT_PATTERNS = [
    r'единица\s*(измерения)?',
    r'ед\.?\s*изм\.?',
    r'размерность',
]


def _match_any_pattern(text: str, patterns: List[str]) -> bool:
    t = text.lower().strip()
    return any(re.search(p, t) for p in patterns)


def _detect_characteristics_columns(table) -> Optional[Dict[str, Any]]:
    col_map: Dict[str, Optional[int]] = {
        "item_name": None,
        "item_number": None,
        "char_name": None,
        "value": None,
        "unit": None,
    }
    header_rows = 1

    for row_idx in range(min(3, len(table.rows))):
        cells = [cell.text.strip() for cell in table.rows[row_idx].cells]

        for col_idx, cell_text in enumerate(cells):
            if not cell_text:
                continue

            if col_map["item_name"] is None and _match_any_pattern(cell_text, _ITEM_NAME_PATTERNS):
                col_map["item_name"] = col_idx
                header_rows = max(header_rows, row_idx + 1)

            if col_map["item_number"] is None and _match_any_pattern(cell_text, _ITEM_NUMBER_PATTERNS):
                col_map["item_number"] = col_idx
                header_rows = max(header_rows, row_idx + 1)

            if col_map["char_name"] is None and _match_any_pattern(cell_text, _CHAR_NAME_PATTERNS):
                col_map["char_name"] = col_idx
                header_rows = max(header_rows, row_idx + 1)

            if col_map["value"] is None and _match_any_pattern(cell_text, _VALUE_PATTERNS):
                col_map["value"] = col_idx
                header_rows = max(header_rows, row_idx + 1)

            if col_map["unit"] is None and _match_any_pattern(cell_text, _UNIT_PATTERNS):
                col_map["unit"] = col_idx
                header_rows = max(header_rows, row_idx + 1)

    if col_map["char_name"] is None or col_map["value"] is None:
        if col_map["char_name"] is not None and col_map["value"] is None:
            fallback_value = col_map["char_name"] + 1
            if fallback_value < len(table.columns):
                col_map["value"] = fallback_value
        else:
            return None

    col_map["header_rows"] = header_rows
    return col_map


def _get_cell(cells: List[str], idx: Optional[int]) -> str:
    if idx is None or idx >= len(cells):
        return ""
    return cells[idx]


# ═══════════════════════════════════════════════════════════════════════════
# Equipment List Table

def _extract_equipment_list(table) -> Dict[str, int]:
    if len(table.rows) < 2:
        return {}

    first_row = [cell.text.strip().lower() for cell in table.rows[0].cells]

    has_name = any(
        "наименование" in cell or "оборудование" in cell or "товар" in cell
        for cell in first_row
    )
    has_qty = any(
        "количество" in cell or "кол-во" in cell or "шт" in cell or "qty" in cell
        for cell in first_row
    )

    if not (has_name and has_qty):
        return {}

    name_col: Optional[int] = None
    qty_col: Optional[int] = None
    for idx, cell in enumerate(first_row):
        if name_col is None and ("наименование" in cell or "оборудование" in cell or "товар" in cell):
            name_col = idx
        if qty_col is None and ("количество" in cell or "кол-во" in cell or "шт" in cell or "qty" in cell):
            qty_col = idx

    if name_col is None:
        return {}

    result: Dict[str, int] = {}
    for row in table.rows[1:]:
        cells = [cell.text.strip() for cell in row.cells]
        if not any(cells):
            continue

        name = _get_cell(cells, name_col)
        if not name:
            continue

        qty = 1
        if qty_col is not None:
            qty_str = _get_cell(cells, qty_col)
            digits = re.findall(r'\d+', qty_str)
            if digits:
                qty = int(digits[0])

        result[name.lower()] = qty

    return result


def _match_quantity(item_name: str, equipment_list: Dict[str, int]) -> Optional[int]:
    name_lower = item_name.lower().strip()
    for list_name, qty in equipment_list.items():
        if name_lower in list_name or list_name in name_lower:
            return qty
    return None



def _parse_table_rows(table, col_map: Dict[str, Any]) -> List[Dict[str, Any]]:
    """Разбор строк таблицы характеристик. char_name — оригинальное название без нормализации."""
    header_rows: int = col_map.get("header_rows", 1)
    parsed_rows = []

    last_item_name = ""
    last_item_number = ""

    for row in table.rows[header_rows:]:
        cells = [cell.text.strip() for cell in row.cells]

        if not any(cells):
            continue

        item_name = _get_cell(cells, col_map["item_name"]) or last_item_name
        item_number = _get_cell(cells, col_map["item_number"]) or last_item_number
        char_name = _get_cell(cells, col_map["char_name"])
        value = _get_cell(cells, col_map["value"])
        unit = _get_cell(cells, col_map["unit"])

        if _get_cell(cells, col_map["item_name"]):
            last_item_name = item_name
        if _get_cell(cells, col_map["item_number"]):
            last_item_number = item_number

        if not char_name:
            continue

        # Пропускаем строки-подзаголовки внутри таблицы
        if _match_any_pattern(char_name, _CHAR_NAME_PATTERNS):
            continue

        char_name = char_name.replace('\xa0', ' ').strip()

        value_with_unit = f"{value} {unit}".strip() if unit and unit.strip() else value
        parsed_value = parse_value(value_with_unit)

        parsed_rows.append({
            "item_name": item_name,
            "item_number": item_number,
            "char_name": char_name,
            "value": value,
            "unit": unit,
            "parsed_value": parsed_value,
        })

    return parsed_rows


def _group_requirements_by_item(rows: List[Dict[str, Any]]) -> Dict[str, List[Dict[str, Any]]]:
    groups: Dict[str, List[Dict[str, Any]]] = defaultdict(list)

    for row in rows:
        item_num = row["item_number"]
        match = re.match(r'^(\d+)', item_num)
        if match:
            prefix = match.group(1)
        elif row["item_name"]:
            prefix = row["item_name"]
        else:
            prefix = "default"

        groups[prefix].append(row)

    return groups


def _extract_model_name_from_text(text: str) -> Optional[str]:
    """Извлечь название модели типа 'MES2300DI-28' из строки 'MES2300DI-28 или эквивалент'."""
    m = re.match(r'^([A-Z][A-Z0-9][-A-Z0-9/.]+)', text.strip(), re.IGNORECASE)
    if m:
        name = m.group(1).strip()
        if re.search(r'\d', name):  # название должно содержать хотя бы одну цифру
            return name
    return None


def _build_item_dict(
    item_prefix: str,
    requirements: List[Dict[str, Any]],
    equipment_list: Dict[str, int],
) -> Optional[Dict[str, Any]]:
    if not requirements:
        return None

    item_names = [r["item_name"] for r in requirements if r["item_name"]]
    item_name = item_names[0] if item_names else f"Позиция {item_prefix}"

    # Определяем категорию по названию позиции
    category = None
    item_lower = item_name.lower()
    if "коммутатор" in item_lower or "switch" in item_lower:
        category = "Коммутаторы"
    elif "маршрутизатор" in item_lower or "router" in item_lower or "шлюз" in item_lower:
        category = "Маршрутизаторы"

    # Пробуем извлечь название модели из строки, например "MES2300DI-28 или эквивалент"
    model_name = _extract_model_name_from_text(item_name)

    required_specs: Dict[str, Any] = {}
    for req in requirements:
        char_name = req["char_name"]
        parsed_value = req["parsed_value"]
        if parsed_value is not None and char_name:
            if char_name not in required_specs:  # берём первое вхождение
                required_specs[char_name] = parsed_value

    quantity = _match_quantity(item_name, equipment_list)

    return {
        "item_name": f"{item_name} (позиция {item_prefix})" if item_prefix.isdigit() else item_name,
        "quantity": quantity,
        "model_name": model_name,
        "category": category,
        "required_specs": required_specs,
    }



def parse_requirements_from_tables(file_path: str) -> Optional[Dict[str, Any]]:
    """
    Разбор требований из структурированных таблиц DOCX.
    Возвращает None, если подходящих таблиц не найдено.
    """
    try:
        doc = Document(file_path)
    except Exception as e:
        logger.error(f"Не удалось открыть DOCX: {e}")
        return None

    logger.info(f"Анализ {len(doc.tables)} таблиц документа")

    characteristics_tables: List[Tuple[int, Any, Dict]] = []
    equipment_list: Dict[str, int] = {}

    for idx, table in enumerate(doc.tables):
        col_map = _detect_characteristics_columns(table)
        if col_map is not None:
            logger.info(f"Таблица характеристик: индекс {idx} ({len(table.rows)} строк)")
            characteristics_tables.append((idx, table, col_map))
        else:
            eq_list = _extract_equipment_list(table)
            if eq_list:
                logger.info(f"Таблица перечня оборудования: индекс {idx} ({len(eq_list)} позиций)")
                equipment_list.update(eq_list)

    if not characteristics_tables:
        logger.info("Структурированная таблица характеристик не найдена")
        return None

    all_parsed_rows: List[Dict[str, Any]] = []
    for idx, table, col_map in characteristics_tables:
        rows = _parse_table_rows(table, col_map)
        logger.info(f"Таблица {idx}: разобрано {len(rows)} строк требований")
        all_parsed_rows.extend(rows)

    if not all_parsed_rows:
        return None

    groups = _group_requirements_by_item(all_parsed_rows)
    logger.info(f"Сгруппировано в {len(groups)} позиций оборудования")

    items = []
    for prefix in sorted(groups.keys(), key=lambda x: int(x) if x.isdigit() else 0):
        item_dict = _build_item_dict(prefix, groups[prefix], equipment_list)
        if item_dict:
            items.append(item_dict)

    result = {"items": items}
    logger.info(
        f"Table parser: {len(items)} позиций, "
        f"всего характеристик: {sum(len(item['required_specs']) for item in items)}"
    )
    return result



def _parse_inline_block(text: str) -> Dict[str, Any]:
    """Разбор блока, разделённого точкой с запятой: «Питание: 100–240 В; Кол-во: 28; ...»"""
    specs: Dict[str, Any] = {}
    for part in re.split(r';\s*', text):
        part = part.strip()
        if not part:
            continue
        m = re.match(r'^([^:]+):\s*(.+)$', part)
        if not m:
            continue
        key = m.group(1).strip().replace('\xa0', ' ')
        value_raw = m.group(2).strip()
        parsed = parse_value(value_raw)
        if parsed is not None:
            specs[key] = parsed
    return specs


def parse_inline_descriptions(file_path: str) -> Optional[Dict[str, Any]]:
    """
    Разбор строчного формата из DOCX.
    Поддерживает пары «Характеристика: значение» — по одной на строку или через «;».
    Новые позиции разделяются абзацем «Наименование товара:».
    """
    try:
        doc = Document(file_path)
    except Exception as e:
        logger.error(f"Не удалось открыть DOCX для inline-парсинга: {e}")
        return None

    items: List[Dict[str, Any]] = []
    current_specs: Dict[str, Any] = {}
    current_item_name: Optional[str] = None
    current_model_name: Optional[str] = None
    current_category: Optional[str] = None

    def _flush_item():
        nonlocal current_specs, current_item_name, current_model_name, current_category
        if current_specs:
            # Определяем категорию по названию, если не задана явно
            cat = current_category
            if cat is None and current_item_name:
                il = current_item_name.lower()
                if "коммутатор" in il or "switch" in il:
                    cat = "Коммутаторы"
                elif "маршрутизатор" in il or "router" in il or "шлюз" in il or "esr" in il:
                    cat = "Маршрутизаторы"

            items.append({
                "item_name": current_item_name or f"Позиция {len(items) + 1}",
                "quantity": None,
                "model_name": current_model_name,
                "category": cat,
                "required_specs": dict(current_specs),
            })
        current_specs = {}
        current_item_name = None
        current_model_name = None
        current_category = None

    # Собираем текст из абзацев и ячеек таблиц
    text_lines: List[str] = []
    for para in doc.paragraphs:
        t = para.text.replace('\xa0', ' ').strip()
        if t:
            text_lines.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.replace('\xa0', ' ').strip()
                if t:
                    text_lines.append(t)

    for line in text_lines:
        m = re.match(r'^([^:]{2,60}):\s*(.+)$', line, re.DOTALL)
        if not m:
            if ';' in line:
                block_specs = _parse_inline_block(line)
                current_specs.update(block_specs)
            continue

        key = m.group(1).strip()
        value_raw = m.group(2).strip()
        key_lower = key.lower()

        # «Наименование товара» — начало новой позиции
        if re.search(r'наименование\s*(товара|изделия|позиции|оборудования)', key_lower):
            _flush_item()
            current_item_name = value_raw
            current_model_name = _extract_model_name_from_text(value_raw)
            continue

        # Определение категории
        if re.search(r'(категория|тип\s*устройства|тип\s*оборудования)', key_lower):
            vl = value_raw.lower()
            if "коммутатор" in vl or "switch" in vl:
                current_category = "Коммутаторы"
            elif "маршрутизатор" in vl or "router" in vl or "шлюз" in vl:
                current_category = "Маршрутизаторы"
            continue

        # Количество — не характеристика, пропускаем
        if re.search(r'количество\s*(единиц|шт\.?|штук)', key_lower):
            continue

        # Если в значении есть «;» — несколько характеристик на одной строке
        if ';' in value_raw:
            parsed = parse_value(value_raw.split(';')[0])
            if parsed is not None:
                current_specs[key] = parsed
            rest = ';'.join(value_raw.split(';')[1:])
            extra = _parse_inline_block(rest)
            current_specs.update(extra)
        else:
            parsed = parse_value(value_raw)
            if parsed is not None:
                current_specs[key] = parsed

    _flush_item()

    if not items:
        logger.info("Inline parser: позиции не найдены")
        return None

    logger.info(
        f"Inline parser: {len(items)} позиций, "
        f"всего характеристик: {sum(len(it['required_specs']) for it in items)}"
    )
    return {"items": items}
