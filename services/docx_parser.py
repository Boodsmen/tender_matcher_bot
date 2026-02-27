"""Извлечение текста из DOCX-файлов (абзацы + таблицы)."""

from docx import Document
from utils.logger import logger


def extract_text_from_docx(file_path: str) -> str:
    """
    Извлечь весь текст из DOCX, включая абзацы и таблицы.
    Вызывает ValueError, если файл не удаётся разобрать.
    """
    try:
        doc = Document(file_path)
    except Exception as e:
        logger.error(f"Не удалось открыть DOCX {file_path}: {e}")
        raise ValueError(f"Не удалось открыть DOCX файл: {e}") from e

    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        table_lines: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            row_text = " | ".join(cells)
            if row_text.replace("|", "").strip():
                table_lines.append(row_text)
        if table_lines:
            parts.append("\n".join(table_lines))

    full_text = "\n\n".join(parts)
    logger.info(f"Извлечено {len(full_text)} символов из DOCX ({len(doc.paragraphs)} абзацев, {len(doc.tables)} таблиц)")
    return full_text
